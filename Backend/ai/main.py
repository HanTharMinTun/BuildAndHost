from fastapi import FastAPI, File, Form, UploadFile
from typing import List, Optional
import os
import sys
from uuid import uuid4
from pathlib import Path

import website_planner
import theme_designer
import layout_designer
from fastapi import Body

from fastapi.middleware.cors import CORSMiddleware
from fastapi.exceptions import RequestValidationError
from fastapi.responses import JSONResponse
from starlette.requests import Request
from fastapi.staticfiles import StaticFiles

# Add websiteBuilder_Backend to path for document extractors
backend_path = Path(__file__).parent.parent / "websiteBuilder_Backend"
if str(backend_path) not in sys.path:
    sys.path.insert(0, str(backend_path))

# Import document extractors
try:
    from app.document_extractors import process_document, build_ai_context, detect_document_type
    from app.document_extractors.storage import DocumentImageStorage
    DOCUMENT_EXTRACTION_AVAILABLE = True
except ImportError as e:
    print(f"Warning: Document extraction not available: {e}")
    DOCUMENT_EXTRACTION_AVAILABLE = False

app = FastAPI()


@app.exception_handler(RequestValidationError)
async def validation_exception_handler(request: Request, exc: RequestValidationError):
    # Avoid jsonable_encoder attempting to decode binary data from uploaded
    # files when FastAPI's default handler tries to include the raw errors.
    return JSONResponse(
        status_code=422,
        content={"detail": "Request validation failed. Check form fields and file uploads."},
    )

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173"],
    allow_credentials=True,
    allow_headers=["*"],
    allow_methods=["*"],
)


@app.get("/api/ai/")
def read_root():
    return {"message": "Hello User"}


# @app.post("/api/post_prompt")
# async def post_prompt(
#     prompt: str = Form(...),
#     type: str = Form(...),
#     files: Optional[List[UploadFile]] = File(None),
# ):
#     """
#     Accepts a prompt and optional file uploads. Saved files are placed into
#     the frontend `public/uploads` folder so they can be referenced at
#     `/uploads/<filename>` from the dev server.
#     """
#     upload_urls: List[str] = []
#     file_texts = {}
#     if files:
#         uploads_dir = os.path.join(os.path.dirname(__file__), "..", "ai-website-builder", "public", "uploads")
#         uploads_dir = os.path.normpath(uploads_dir)
#         os.makedirs(uploads_dir, exist_ok=True)
#         for f in files:
#             # create a safe unique filename
#             extension = os.path.splitext(f.filename)[1] or ""
#             safe_name = f"{uuid4().hex}{extension}"
#             dest_path = os.path.join(uploads_dir, safe_name)
#             with open(dest_path, "wb") as out:
#                 content = await f.read()
#                 out.write(content)
#             # Use a relative URL the frontend can request from Vite dev server
#             upload_urls.append(f"/uploads/{safe_name}")
#             # Attempt simple text extraction for PDFs and text files.
#             try:
#                 ext = extension.lower()
#                 if ext == ".pdf":
#                     try:
#                         from PyPDF2 import PdfReader
#                         reader = PdfReader(dest_path)
#                         text = []
#                         for page in reader.pages:
#                             page_text = page.extract_text() or ""
#                             text.append(page_text)
#                         file_texts[f"/uploads/{safe_name}"] = "\n".join(text)
#                     except Exception:
#                         # PDF extraction not available or failed; skip
#                         pass
#                 elif ext == ".docx":
#                     try:
#                         from docx import Document
#                         doc = Document(dest_path)
#                         paragraphs = [p.text for p in doc.paragraphs if p.text]
#                         file_texts[f"/uploads/{safe_name}"] = "\n".join(paragraphs)
#                     except Exception:
#                         pass
#                 elif ext in (".txt",):
#                     try:
#                         with open(dest_path, "r", encoding="utf-8", errors="ignore") as fh:
#                             file_texts[f"/uploads/{safe_name}"] = fh.read()
#                     except Exception:
#                         pass
#             except Exception:
#                 pass

#     # First pass creates structure/content only. The second pass creates the
#     # theme for that exact component tree. Pass the upload URLs to the planner
#     # so it can reference attachments in generated Image/CDNIcon props.
#     try:
#         website = website_planner.create_layout(prompt, file_urls=upload_urls, file_texts=file_texts)
#         # Generate responsive layout information for the component tree.
#         try:
#             layout_info = layout_designer.create_responsive_layout(website)
#         except Exception:
#             layout_info = None
#         theme = theme_designer.create_theme(prompt, website)
#         return {"website": website, "layout": layout_info, "theme": theme, "uploads": upload_urls}
#     except Exception as e:
#         # Return a safe JSON response instead of 500 to the client with a helpful message.
#         return JSONResponse(status_code=400, content={"detail": "Planner failed to generate JSON. Try providing key text info or upload plain-text files. Error: " + str(e)})


from fastapi import UploadFile, File, Form
from fastapi.responses import JSONResponse
from typing import Optional, List
from uuid import uuid4
import os


# Public upload directory served by Nginx
UPLOAD_DIR = "/var/www/onlinegif/uploads"


@app.post("/api/ai/post_prompt")
async def post_prompt(
    prompt: str = Form(...),
    type: str = Form(...),
    files: Optional[List[UploadFile]] = File(None),
):
    """
    Accepts a prompt and optional file uploads.

    Uploaded files are stored in:
        /var/www/onlinegif/uploads/

    Nginx serves them as:
        https://webcreator.site/uploads/<filename>
    """

    upload_urls: List[str] = []
    file_texts = {}
    image_urls: List[str] = []  # Track images separately for multimodal Claude input
    
    # Constants
    MAX_FILES = 5

    # Make sure upload directory exists
    os.makedirs(UPLOAD_DIR, exist_ok=True)

    # Validate file count
    if files and len(files) > MAX_FILES:
        return JSONResponse(
            status_code=400,
            content={
                "detail": f"Too many files. Maximum {MAX_FILES} files allowed per message."
            }
        )

    if files:
        for f in files:

            # Get extension safely
            original_filename = f.filename or ""
            extension = os.path.splitext(original_filename)[1].lower()

            # Generate unique filename
            safe_name = f"{uuid4().hex}{extension}"

            # Absolute filesystem path
            dest_path = os.path.join(
                UPLOAD_DIR,
                safe_name
            )

            # Save uploaded file
            with open(dest_path, "wb") as out:
                content = await f.read()
                out.write(content)

            # Public URL - use absolute URL for cross-domain compatibility
            upload_url = f"https://webcreator.site/uploads/{safe_name}"
            upload_urls.append(upload_url)
            
            # Track if this is an image file for Claude vision
            image_extensions = {".jpg", ".jpeg", ".png", ".gif", ".webp"}
            if extension in image_extensions:
                image_urls.append(upload_url)

            # ------------------------------------------------
            # Extract text and images from uploaded files
            # ------------------------------------------------

            try:
                # Detect document type
                doc_type = detect_document_type(dest_path) if DOCUMENT_EXTRACTION_AVAILABLE else None
                
                # Use enhanced extraction for PDF/DOCX (extracts text + images)
                if doc_type in ["pdf", "docx"] and DOCUMENT_EXTRACTION_AVAILABLE:
                    try:
                        print(f"Processing {doc_type.upper()} with enhanced extraction: {original_filename}")
                        
                        # Configure storage to use the AI upload directory
                        storage_handler = DocumentImageStorage(base_upload_dir=UPLOAD_DIR)
                        
                        # Generate document ID
                        doc_id = uuid4().hex
                        
                        # Process document (extracts text AND images)
                        document_data = process_document(
                            file_path=dest_path,
                            document_id=doc_id,
                            storage_handler=storage_handler
                        )
                        
                        # Add extracted text for AI context
                        extracted_text = document_data.get("extracted_text", "")
                        if extracted_text:
                            file_texts[upload_url] = extracted_text
                        
                        # Add extracted images to upload_urls and image_urls for AI vision
                        extracted_images = document_data.get("image_urls", [])
                        for img_url in extracted_images:
                            upload_urls.append(img_url)
                            image_urls.append(img_url)
                        
                        print(f"Enhanced extraction complete: {len(extracted_text)} chars, {len(extracted_images)} images")
                        
                    except Exception as e:
                        print(f"Enhanced {doc_type.upper()} extraction failed: {e}")
                        # Fall back to basic extraction
                        if extension == ".pdf":
                            try:
                                from PyPDF2 import PdfReader
                                reader = PdfReader(dest_path)
                                text = [page.extract_text() or "" for page in reader.pages]
                                file_texts[upload_url] = "\n".join(text)
                                print(f"Fell back to basic PDF extraction")
                            except Exception as e2:
                                print(f"Basic PDF extraction also failed: {e2}")
                        
                        elif extension == ".docx":
                            try:
                                from docx import Document
                                doc = Document(dest_path)
                                paragraphs = [p.text for p in doc.paragraphs if p.text]
                                file_texts[upload_url] = "\n".join(paragraphs)
                                print(f"Fell back to basic DOCX extraction")
                            except Exception as e2:
                                print(f"Basic DOCX extraction also failed: {e2}")
                
                elif extension == ".txt":
                    # Handle text files
                    try:
                        with open(dest_path, "r", encoding="utf-8", errors="ignore") as fh:
                            file_texts[upload_url] = fh.read()
                        print(f"Extracted text file: {original_filename}")
                    except Exception as e:
                        print(f"TXT extraction failed: {e}")
                
                else:
                    # Unsupported file type or extraction not available
                    # Try basic extraction if it's PDF/DOCX
                    if extension == ".pdf":
                        try:
                            from PyPDF2 import PdfReader
                            reader = PdfReader(dest_path)
                            text = [page.extract_text() or "" for page in reader.pages]
                            file_texts[upload_url] = "\n".join(text)
                            print(f"Basic PDF extraction: {original_filename}")
                        except Exception as e:
                            print(f"PDF extraction failed: {e}")
                    
                    elif extension == ".docx":
                        try:
                            from docx import Document
                            doc = Document(dest_path)
                            paragraphs = [p.text for p in doc.paragraphs if p.text]
                            file_texts[upload_url] = "\n".join(paragraphs)
                            print(f"Basic DOCX extraction: {original_filename}")
                        except Exception as e:
                            print(f"DOCX extraction failed: {e}")

            except Exception as e:
                print(f"File processing failed: {e}")

    # ------------------------------------------------
    # AI Website Generation
    # ------------------------------------------------

    try:

        website = website_planner.create_layout(
            prompt,
            file_urls=upload_urls,
            file_texts=file_texts,
            image_urls=image_urls  # Pass images separately for multimodal input
        )

        # Generate responsive layout
        try:
            layout_info = (
                layout_designer.create_responsive_layout(
                    website
                )
            )

        except Exception as e:
            print(f"Layout generation failed: {e}")
            layout_info = None

        # Generate theme
        theme = theme_designer.create_theme(
            prompt,
            website
        )

        return {
            "website": website,
            "layout": layout_info,
            "theme": theme,
            "uploads": upload_urls
        }

    except Exception as e:

        print(f"Website generation failed: {e}")

        return JSONResponse(
            status_code=400,
            content={
                "detail": (
                    "Planner failed to generate JSON. "
                    "Try providing key text information "
                    "or uploading plain-text files. "
                    f"Error: {str(e)}"
                )
            }
        )

        


@app.post("/api/design_layout")
async def design_layout(body: dict = Body(...)):
    """Accept a website component tree (JSON) and return responsive layout info."""
    try:
        layout = layout_designer.create_responsive_layout(body)
        return {"layout": layout}
    except Exception as e:
        return JSONResponse(status_code=400, content={"detail": "Layout designer failed: " + str(e)})


@app.post("/api/ai/design_theme")
async def design_theme(body: dict = Body(...)):
    """Accept a website component tree (JSON) and return a generated theme."""
    try:
        theme = theme_designer.create_theme(body.get("brief", ""), body)
        return {"theme": theme}
    except Exception as e:
        return JSONResponse(status_code=400, content={"detail": "Theme designer failed: " + str(e)})
