"""
DOCX document extraction module.

Extracts text and images from Microsoft Word (.docx) files using python-docx.
Handles paragraphs, headings, and embedded images.
"""

import logging
from pathlib import Path
from typing import Dict, List, Any
import uuid

logger = logging.getLogger(__name__)


def extract_docx_content(
    docx_path: str,
    document_id: str = None,
    storage_handler=None
) -> Dict[str, Any]:
    """
    Extract text and images from a DOCX file.
    
    Args:
        docx_path: Path to the DOCX file
        document_id: Unique identifier for the document (auto-generated if None)
        storage_handler: DocumentImageStorage instance for saving images
        
    Returns:
        Dictionary containing extracted text and image URLs
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
    except ImportError:
        logger.error("python-docx not installed")
        raise ImportError("python-docx is required for DOCX extraction")
    
    if document_id is None:
        document_id = str(uuid.uuid4())
    
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {docx_path}")
    
    logger.info(f"Processing DOCX: {docx_path.name}")
    
    # Initialize storage if not provided
    if storage_handler is None:
        from .storage import DocumentImageStorage
        storage_handler = DocumentImageStorage()
    
    result = {
        "source_file": docx_path.name,
        "source_type": "docx",
        "document_id": document_id,
        "content": [],
        "extracted_text": "",
        "image_urls": []
    }
    
    try:
        # Open document
        doc = Document(str(docx_path))
        
        all_text_parts = []
        total_images = 0
        position = 0
        
        # Extract images from document relationships
        image_parts = {}
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_parts[rel.rId] = rel.target_part
        
        logger.info(f"Found {len(image_parts)} image parts")
        
        # Process paragraphs and embedded content
        for para_idx, paragraph in enumerate(doc.paragraphs):
            position += 1
            
            # Extract text from paragraph
            para_text = paragraph.text.strip()
            
            if para_text:
                result["content"].append({
                    "type": "text",
                    "position": position,
                    "content": para_text
                })
                all_text_parts.append(para_text)
                logger.debug(f"Position {position}: Extracted {len(para_text)} characters")
            
            # Check for inline images in this paragraph
            for run in paragraph.runs:
                if hasattr(run, '_element'):
                    for drawing in run._element.findall(f".//{qn('w:drawing')}"):
                        for inline in drawing.findall(f".//{qn('pic:pic')}"):
                            try:
                                # Extract image reference
                                blip = inline.find(f".//{qn('a:blip')}")
                                if blip is not None:
                                    embed_id = blip.get(qn('r:embed'))
                                    if embed_id and embed_id in image_parts:
                                        image_part = image_parts[embed_id]
                                        image_bytes = image_part.blob
                                        
                                        # Determine extension
                                        content_type = image_part.content_type
                                        ext_map = {
                                            "image/png": "png",
                                            "image/jpeg": "jpg",
                                            "image/jpg": "jpg",
                                            "image/gif": "gif",
                                            "image/webp": "webp"
                                        }
                                        extension = ext_map.get(content_type, "png")
                                        
                                        # Save image
                                        _, image_url = storage_handler.save_image(
                                            image_data=image_bytes,
                                            document_id=document_id,
                                            image_index=total_images,
                                            extension=extension
                                        )
                                        
                                        # Add image content
                                        position += 1
                                        result["content"].append({
                                            "type": "image",
                                            "position": position,
                                            "url": image_url,
                                            "index": total_images
                                        })
                                        result["image_urls"].append(image_url)
                                        
                                        total_images += 1
                                        logger.debug(f"Position {position}: Extracted image {total_images}")
                                        
                            except Exception as e:
                                logger.warning(f"Failed to extract inline image: {e}")
                                continue
        
        # Consolidate all text
        result["extracted_text"] = "\n\n".join(all_text_parts)
        result["total_paragraphs"] = len(doc.paragraphs)
        
        logger.info(
            f"DOCX processing complete: "
            f"{len(doc.paragraphs)} paragraphs, "
            f"{len(result['extracted_text'])} chars, "
            f"{total_images} images"
        )
        
        return result
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {e}", exc_info=True)
        raise


def get_docx_metadata(docx_path: str) -> Dict[str, Any]:
    """
    Extract basic metadata from a DOCX file without processing content.
    
    Args:
        docx_path: Path to the DOCX file
        
    Returns:
        Dictionary with metadata (paragraph count, author, etc.)
    """
    try:
        from docx import Document
    except ImportError:
        return {"error": "python-docx not installed"}
    
    try:
        doc = Document(str(docx_path))
        metadata = {
            "paragraph_count": len(doc.paragraphs),
            "author": doc.core_properties.author,
            "title": doc.core_properties.title,
            "subject": doc.core_properties.subject,
        }
        return metadata
    except Exception as e:
        logger.error(f"Failed to read DOCX metadata: {e}")
        return {"error": str(e)}

